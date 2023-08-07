from docx import Document
from docx.shared import RGBColor

import re
import html
from urllib import parse
import requests

import pandas as pd

GOOGLE_TRANSLATE_URL = 'http://translate.google.com/m?q=%s&tl=%s&sl=%s'


def get_red_words(doc_path):
    doc = Document(doc_path)
    red_words = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color.rgb == RGBColor(255, 0, 0):
                print("检测到生词:", run.text.split()[0])
                red_words.extend(run.text.split())
    return red_words


def translate(text, to_language="zh-CN", text_language="en"):
    text = parse.quote(text)
    url = GOOGLE_TRANSLATE_URL % (text, to_language, text_language)
    response = requests.get(url)
    data = response.text
    expr = r'(?s)class="(?:t0|result-container)">(.*?)<'
    result = re.findall(expr, data)
    if (len(result) == 0):
        return ""
    return html.unescape(result[0])


# Define the function to add mapped string to red words
def add_mapped_string_to_red_words(paragraph):
    for run in paragraph.runs:
        if run.font.color and isinstance(run.font.color.rgb, RGBColor):
            color_value = run.font.color.rgb
            if color_value == RGBColor(255, 0, 0):  # Check if the color is red
                # Check if the word exists in the mapping dictionary
                if run.text.lower() in translations:
                    run.text = run.text + "(" + translations[run.text.lower()] + ")"


if __name__ == "__main__":

    red_words = get_red_words("input.docx")

    # 用于存放翻译结果的字典
    translations = {}

    for word in red_words:
        # 调用 translate 方法进行翻译，源语言为英文('en')，目标语言为中文('zh-cn')
        translation = translate(word)

        print("翻译成功:", word, translation)
        # 将结果保存在字典中
        translations[word] = translation

    doc = Document("input.docx")

    # Load the data from the Excel file
    df1 = pd.read_excel("TOEFL.xlsx")

    # Extract the words into a list
    words_list = df1.iloc[:, 0].str.lower().tolist()

    TOEFL = {}

    # For each paragraph in the document
    for para in doc.paragraphs:
        # For each run in the paragraph
        for run in para.runs:
            # Check each word in the run
            for word in run.text.split():
                # If the word (in lowercase) is in the list of words
                if word.lower() in words_list:
                    TOEFL[word.lower()] = df1[df1.Word == word.lower()]['Chinese'].values[0]

    # Load the data from the Excel file
    df2 = pd.read_excel("GRE.xlsx")

    # Extract the words into a list
    words_list = df2.iloc[:, 0].str.lower().tolist()

    GRE = {}

    # For each paragraph in the document
    for para in doc.paragraphs:
        # For each run in the paragraph
        for run in para.runs:
            # Only consider runs that are not red
            for word in run.text.split():
                # If the word (in lowercase) is in the list of words
                if word.lower() in words_list:
                    GRE[word.lower()] = df2[df2.Word == word.lower()]['Chinese'].values[0]

    for paragraph in doc.paragraphs:
        add_mapped_string_to_red_words(paragraph)

    doc.save("output.docx")

    # 打开一个文件，如果不存在则创建
    with open('output.md', 'w', encoding='utf-8') as f:

        f.write('### 文内生词\n')
        # 写入markdown表格头
        f.write('| English | Chinese |\n')
        f.write('|---------|---------|\n')

        # 遍历字典，写入每一行内容
        for key, value in translations.items():
            f.write(f'| {key} | {value} |\n')

        f.write('\n### TOEFL单词\n')
        # 写入markdown表格头
        f.write('| English | Chinese |\n')
        f.write('|---------|---------|\n')

        # 遍历字典，写入每一行内容
        for key, value in TOEFL.items():
            f.write(f'| {key} | {value} |\n')

        f.write('\n### GRE单词\n')
        # 写入markdown表格头
        f.write('| English | Chinese |\n')
        f.write('|---------|---------|\n')

        # 遍历字典，写入每一行内容
        for key, value in GRE.items():
            f.write(f'| {key} | {value} |\n')

    print("数据已经成功导出！")
